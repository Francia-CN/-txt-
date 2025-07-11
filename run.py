# -*- coding: utf-8 -*-
"""
多格式文档批量转换txt工具 v1.0

一个高效、健壮、可扩展的命令行工具，用于将多种格式的文档（如 EPUB, MOBI,
PDF, DOCX）批量转换为纯文本（.txt）文件。
"""

# --- 1. 标准库导入 ---
import os
import shutil
import zipfile
import argparse
import tempfile
import urllib.parse
import sys
import concurrent.futures
import time
import logging
import warnings
import subprocess
from collections import defaultdict
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Tuple, Any, Literal, Callable

# --- 2. 第三方库导入 (带可选依赖处理) ---
try:
    from bs4 import BeautifulSoup, XMLParsedAsHTMLWarning
    import lxml.etree as etree
    from tqdm import tqdm
    HAVE_LXML_BS4 = True
except ImportError:
    HAVE_LXML_BS4 = False

try:
    import docx
    HAVE_DOCX = True
except ImportError:
    HAVE_DOCX = False

try:
    import fitz  # PyMuPDF
    HAVE_PDF = True
except ImportError:
    HAVE_PDF = False

try:
    if shutil.which("ebook-convert"):
        HAVE_CALIBRE = True
    else:
        HAVE_CALIBRE = False
except Exception:
    HAVE_CALIBRE = False

# --- 3. 全局常量与配置 ---
VERSION = "v1.0"
LOG_DIR_NAME = "logs"
DEFAULT_TIMEOUT = 120

EPUB_TAGS_TO_REMOVE = [
    "script", "style", "nav", "header", "footer", "meta", "link",
    "noscript", "svg", "figure", "figcaption", "a", "img", "aside"
]

# --- 4. 自定义类型别名 ---
HandlerResult = Tuple[str, List[str]]
HandlerFunction = Callable[[Path], HandlerResult]
ProcessResult = Tuple[Literal['converted', 'copied', 'error', 'skipped'], str, int, List[str]]


# --- 5. 辅助类与函数 ---
class Color:
    HEADER = "\033[95m"; OKBLUE = "\033[94m"; OKGREEN = "\033[92m"; WARNING = "\033[93m"
    FAIL = "\033[91m"; ENDC = "\033[0m"; BOLD = "\033[1m"; CYAN = "\033[96m"

def setup_logging(log_dir: Path, level: str) -> Path:
    log_dir.mkdir(exist_ok=True)
    log_file = log_dir / f"conversion_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log"
    logger = logging.getLogger()
    if logger.hasHandlers():
        logger.handlers.clear()
    logger.setLevel(logging.DEBUG)
    file_handler = logging.FileHandler(log_file, encoding="utf-8")
    log_level = getattr(logging, level.upper(), logging.INFO)
    file_handler.setLevel(log_level)
    formatter = logging.Formatter("%(asctime)s [%(levelname)s] (%(processName)s) %(message)s")
    file_handler.setFormatter(formatter)
    logger.addHandler(file_handler)
    return log_file

def format_size(size_bytes: float) -> str:
    if size_bytes < 1024.0:
        return f"{size_bytes:.1f} B"
    for unit in ["B", "KB", "MB", "GB", "TB"]:
        if size_bytes < 1024.0:
            return f"{size_bytes:.1f} {unit}"
        size_bytes /= 1024.0
    return f"{size_bytes:.1f} PB"


# --- 6. 文件格式处理器 (策略模式实现) ---

def handle_epub(src_path: Path) -> HandlerResult:
    warnings_list = []
    with zipfile.ZipFile(src_path, "r") as zf:
        if any(f.flag_bits & 0x1 for f in zf.infolist()):
            raise RuntimeError("文件已加密")
        opf_path = ""
        try:
            container_xml = zf.read("META-INF/container.xml")
            container = etree.fromstring(container_xml)
            ns = {"container": "urn:oasis:names:tc:opendocument:xmlns:container"}
            opf_path_results = container.xpath("//container:rootfile/@full-path", namespaces=ns)
            if opf_path_results: opf_path = opf_path_results[0]
        except (etree.ParseError, KeyError, zipfile.BadZipFile): pass
        if not opf_path:
            opf_candidates = [p for p in zf.namelist() if p.lower().endswith('.opf')]
            if not opf_candidates: raise RuntimeError("在EPUB中找不到OPF配置文件 (.opf)")
            opf_path = opf_candidates[0]
        opf_tree = etree.fromstring(zf.read(opf_path))
        ns_map = {k: v for k, v in opf_tree.nsmap.items() if k} or {'opf': 'http://www.idpf.org/2007/opf'}
        manifest = {item.get("id"): item.get("href") for item in opf_tree.xpath(".//opf:manifest/opf:item", namespaces=ns_map)}
        spine_idrefs = [item.get("idref") for item in opf_tree.xpath(".//opf:spine/opf:itemref", namespaces=ns_map)]
        opf_dir = Path(opf_path).parent
        full_text_parts = []
        for idref in spine_idrefs:
            href = manifest.get(idref)
            if not href: continue
            content_path_str = str(Path(opf_dir) / urllib.parse.unquote(href)).replace("\\", "/")
            if content_path_str in zf.namelist():
                try:
                    content_bytes = zf.read(content_path_str)
                    soup = None
                    try: soup = BeautifulSoup(content_bytes, "lxml-xml")
                    except etree.XMLSyntaxError:
                        with warnings.catch_warnings():
                            warnings.simplefilter("ignore", category=XMLParsedAsHTMLWarning)
                            msg = f"文件 '{content_path_str}' XML格式不规范，已回退到HTML解析器。"
                            warnings_list.append(msg)
                            logging.warning(f"[{src_path.name}] {msg}")
                            soup = BeautifulSoup(content_bytes, "lxml")
                    for tag in soup.find_all(EPUB_TAGS_TO_REMOVE): tag.decompose()
                    body = soup.find("body")
                    if body:
                        text = body.get_text(separator='\n', strip=True)
                        if text: full_text_parts.append(text)
                except Exception as e: logging.debug(f"处理EPUB内文件'{content_path_str}'失败: {e}")
            else: warnings_list.append(f"引用的章节文件'{content_path_str}'未在压缩包中找到")
    return "\n\n".join(full_text_parts), warnings_list

def handle_docx(src_path: Path) -> HandlerResult:
    doc = docx.Document(src_path)
    text_parts = []
    for section in doc.sections:
        for part in (section.header, section.footer):
            for p in part.paragraphs:
                if p.text.strip(): text_parts.append(p.text)
    for p in doc.paragraphs: text_parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text_parts.append(cell.text.strip())
    return "\n".join(filter(None, text_parts)), []

def handle_pdf(src_path: Path) -> HandlerResult:
    """PDF 文件处理器，使用 PyMuPDF 提取文本。"""
    with fitz.open(src_path) as doc:
        if doc.is_encrypted:
            raise RuntimeError("PDF文件已加密")
        # ======================================================================
        #  优化点: 使用 "text" 参数可以更快地提取纯文本，因为它会忽略图片和矢量图形。
        #  这对于将PDF转为纯TXT的场景非常有效。
        # ======================================================================
        return "".join(page.get_text("text") for page in doc), []

def handle_mobi(src_path: Path) -> HandlerResult:
    if not HAVE_CALIBRE: raise RuntimeError("未找到 Calibre 的命令行工具 (ebook-convert)。")
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_txt_path = Path(temp_dir) / f"{src_path.stem}.txt"
        command = ["ebook-convert", str(src_path), str(temp_txt_path)]
        try:
            creation_flags = subprocess.CREATE_NO_WINDOW if sys.platform == "win32" else 0
            process = subprocess.run(command, check=True, capture_output=True, text=True, encoding='utf-8', errors='ignore', creationflags=creation_flags)
            if temp_txt_path.exists():
                text_content = temp_txt_path.read_text(encoding='utf-8')
                if not text_content.strip(): raise RuntimeError("Calibre 转换成功，但输出的文本文件为空。")
                return text_content, []
            else: raise RuntimeError("Calibre 转换过程未生成输出文件。")
        except FileNotFoundError: raise RuntimeError("无法执行 'ebook-convert'。请确保 Calibre 已安装并添加至系统PATH。")
        except subprocess.CalledProcessError as e:
            error_message = (e.stderr or e.stdout or "无详细输出").strip()
            if "DRM" in error_message: raise RuntimeError("MOBI文件可能受DRM保护，Calibre 无法转换。")
            raise RuntimeError(f"Calibre 转换失败。错误: {error_message}")


# --- 7. 核心处理引擎 ---
PROCESSOR_STRATEGIES: Dict[str, HandlerFunction] = {}
if HAVE_LXML_BS4: PROCESSOR_STRATEGIES['.epub'] = handle_epub
if HAVE_DOCX: PROCESSOR_STRATEGIES['.docx'] = handle_docx
if HAVE_PDF: PROCESSOR_STRATEGIES['.pdf'] = handle_pdf
if HAVE_CALIBRE:
    PROCESSOR_STRATEGIES['.mobi'] = handle_mobi
    PROCESSOR_STRATEGIES['.azw3'] = handle_mobi

def process_file(src_path: Path, rel_path: Path, output_dir: Path) -> ProcessResult:
    rel_path_str = str(rel_path)
    ext = src_path.suffix.lower()
    try:
        file_size = src_path.stat().st_size
        if file_size == 0:
            return 'skipped', rel_path_str, 0, ["文件为空(0字节)"]
        processor = PROCESSOR_STRATEGIES.get(ext)
        if processor:
            text, warnings_list = processor(src_path)
            if not text or not text.strip():
                # ======================================================================
                #  优化点: 提供更友好的错误提示，引导用户思考PDF是否为扫描件。
                # ======================================================================
                raise ValueError("提取内容为空 (PDF可能为扫描件或无文本内容)")
            output_path = (output_dir / rel_path).with_suffix(".txt")
            output_path.parent.mkdir(parents=True, exist_ok=True)
            output_path.write_text(text, encoding='utf-8')
            return 'converted', rel_path_str, file_size, warnings_list
        else:
            dest_path = output_dir / rel_path
            dest_path.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(src_path, dest_path)
            return 'copied', rel_path_str, file_size, []
    except Exception as e:
        logging.error(f"处理文件 [{rel_path_str}] 时发生严重错误。", exc_info=True)
        return 'error', rel_path_str, 0, [str(e)]


# --- 8. 主程序与用户界面 ---
def create_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description=f"多格式文档批量转换工具 (版本 {VERSION})", formatter_class=argparse.RawTextHelpFormatter, epilog="""
使用示例:
  1. 基本用法 (自动扫描 'input' 目录, 输出到 'output'):
     python %(prog)s

  2. 指定目录，使用8个进程并强制清空输出目录:
     python %(prog)s ./my_books ./my_texts -t 8 --clean-output

  3. 只转换PDF和EPUB，并以最详细的模式记录日志:
     python %(prog)s ./in ./out --include .pdf,.epub --log-level DEBUG
""")
    parser.add_argument("input", nargs='?', default="input", help="输入目录路径 (默认: 'input')")
    parser.add_argument("output", nargs='?', default="output", help="输出目录路径 (默认: 'output')")
    parser.add_argument("-t", "--threads", type=int, default=os.cpu_count(), help=f"处理进程数 (默认: 系统CPU核心数, {os.cpu_count()})")
    parser.add_argument("--timeout", type=int, default=DEFAULT_TIMEOUT, help=f"单个文件的最大处理时间(秒) (默认: {DEFAULT_TIMEOUT})")
    parser.add_argument("--clean-output", action="store_true", help="在开始前清空已存在的输出目录 (危险操作!)")
    parser.add_argument("--include", type=str, help="只处理指定后缀的文件, 用逗号分隔 (例: .epub,.pdf)")
    parser.add_argument("--exclude", type=str, help="排除指定后缀的文件, 用逗号分隔 (例: .jpg,.png)")
    parser.add_argument("--no-color", action="store_true", help="禁用所有彩色控制台输出")
    parser.add_argument("--log-level", type=str, default="INFO", choices=["DEBUG", "INFO", "WARNING", "ERROR"], help="设置日志文件的记录级别 (默认: INFO)")
    return parser

def print_summary(stats: Dict[str, Any], duration: float, log_file: Path, c: Color):
    print("\n" + "=" * 80)
    print(f"{c.BOLD}处理完成！总耗时: {duration:.2f} 秒{c.ENDC}")
    print(f"{c.OKGREEN}成功转换: {stats['converted']:>5} 文件 ({format_size(stats['converted_size'])}){c.ENDC}")
    print(f"{c.CYAN}直接复制: {stats['copied']:>5} 文件 ({format_size(stats['copied_size'])}){c.ENDC}")
    print(f"{c.OKBLUE}跳过文件: {stats['skipped']:>5} 文件{c.ENDC}")
    if stats['warnings']: print(f"{c.WARNING}产生警告: {stats['warnings']:>5} 文件 (详情请查看日志){c.ENDC}")
    if stats['errors']: print(f"{c.FAIL}发生错误: {stats['errors']:>5} 文件 (详情请查看日志){c.ENDC}")
    if stats['timeouts']: print(f"{c.FAIL}{c.BOLD}处理超时: {stats['timeouts']:>5} 文件{c.ENDC}")
    print(f"\n日志文件已保存至: {c.CYAN}{log_file}{c.ENDC}")
    print("=" * 80)

def main():
    args = create_arg_parser().parse_args()
    c = Color if not args.no_color and sys.stdout.isatty() else type("Color", (), {k: "" for k in Color.__dict__ if not k.startswith('__')})()
    log_file = setup_logging(Path(__file__).parent / LOG_DIR_NAME, args.log_level)
    print(f"\n{c.HEADER}{'='*50}{c.ENDC}")
    print(f"{c.HEADER}    多格式文档批量转换工具 (版本 {VERSION})    {c.ENDC}")
    print(f"{c.HEADER}{'='*50}{c.ENDC}\n")
    if not HAVE_LXML_BS4:
        print(f"{c.FAIL}致命错误: 核心依赖 'lxml', 'beautifulsoup4' 或 'tqdm' 未安装。\n请运行 'pip install lxml beautifulsoup4 tqdm' 后重试。{c.ENDC}")
        sys.exit(1)
    if not PROCESSOR_STRATEGIES:
        print(f"{c.FAIL}致命错误: 未安装任何格式处理器的依赖库 (如 python-docx, PyMuPDF)。\n请根据需要安装，例如 'pip install python-docx PyMuPDF'。{c.ENDC}")
        sys.exit(1)
    for ext, handler_installed in {'.docx': HAVE_DOCX, '.pdf': HAVE_PDF, '.mobi': HAVE_CALIBRE, '.azw3': HAVE_CALIBRE}.items():
        if not handler_installed:
            if ext in ['.mobi', '.azw3']: print(f"{c.WARNING}提示: 未找到 Calibre 环境，将无法转换 {ext} 文件。{c.ENDC}")
            else: print(f"{c.WARNING}提示: 未安装 '{'.docx'=='python-docx' and ext or 'PyMuPDF'}'，将无法转换 {ext} 文件。{c.ENDC}")
    input_dir, output_dir = Path(args.input), Path(args.output)
    stats = defaultdict(int)
    try:
        if not input_dir.is_dir():
            input_dir.mkdir(exist_ok=True)
            print(f"{c.OKBLUE}输入目录 '{input_dir}' 不存在，已自动创建。请放入文件后重新运行。{c.ENDC}")
            sys.exit(0)
        if output_dir.exists():
            if args.clean_output:
                print(f"{c.WARNING}警告: 用户指定了 --clean-output，正在清空输出目录 '{output_dir}'...{c.ENDC}")
                shutil.rmtree(output_dir)
            else:
                print(f"{c.FAIL}错误: 输出目录 '{output_dir}' 已存在。为防止数据丢失，请先手动删除该目录，\n      或使用 --clean-output 标志授权脚本进行清空。{c.ENDC}")
                sys.exit(1)
        output_dir.mkdir(parents=True, exist_ok=True)
        all_files = [f for f in input_dir.rglob("*") if f.is_file()]
        if args.include:
            included = tuple(e.strip().lower() for e in args.include.split(','))
            all_files = [f for f in all_files if f.suffix.lower() in included]
        if args.exclude:
            excluded = tuple(e.strip().lower() for e in args.exclude.split(','))
            all_files = [f for f in all_files if f.suffix.lower() not in excluded]
    except Exception as e:
        logging.critical(f"启动失败: {e}", exc_info=True)
        print(f"{c.FAIL}启动失败: {e}{c.ENDC}")
        sys.exit(1)
    if not all_files:
        print(f"{c.OKBLUE}在输入目录 '{input_dir}' 中未找到任何符合条件的文件。{c.ENDC}")
        sys.exit(0)
    file_queue = [(f, f.relative_to(input_dir)) for f in all_files]
    num_workers = min(args.threads, len(file_queue))
    print(f"发现 {len(file_queue)} 个待处理文件，使用 {num_workers} 个进程开始转换...\n")
    start_time = time.time()
    with concurrent.futures.ProcessPoolExecutor(max_workers=num_workers) as executor:
        future_map = {executor.submit(process_file, f, rel, output_dir): rel for f, rel in file_queue}
        progress_bar = tqdm(total=len(future_map), desc=f"{c.BOLD}处理进度{c.ENDC}", unit="个", ncols=100)
        for future in concurrent.futures.as_completed(future_map):
            rel_path = future_map[future]
            try:
                status, _, size, messages = future.result(timeout=args.timeout)
                stats[status] += 1
                if status == 'converted':
                    stats['converted_size'] += size
                    if not messages: tqdm.write(f"{c.OKGREEN}[已转换] {rel_path}{c.ENDC}")
                    else:
                        stats['warnings'] += 1
                        tqdm.write(f"{c.WARNING}[有警告] {rel_path} (详见日志){c.ENDC}")
                        for msg in messages: logging.warning(f"[{rel_path}] {msg}")
                elif status == 'copied':
                    stats['copied_size'] += size
                    tqdm.write(f"{c.CYAN}[已复制] {rel_path}{c.ENDC}")
                elif status == 'skipped':
                    tqdm.write(f"{c.OKBLUE}[已跳过] {rel_path}: {messages[0]}{c.ENDC}")
                    logging.info(f"[{rel_path}] 跳过: {messages[0]}")
                elif status == 'error':
                    stats['errors'] += 1
                    error_msg_short = messages[0].splitlines()[0]
                    tqdm.write(f"{c.FAIL}[出错了] {rel_path}: {error_msg_short}{c.ENDC}")
            except concurrent.futures.TimeoutError:
                stats['timeouts'] += 1
                tqdm.write(f"{c.FAIL}{c.BOLD}[超  时] {rel_path}: 处理超过 {args.timeout} 秒，已终止{c.ENDC}")
                logging.error(f"处理文件 [{rel_path}] 超时。")
            except Exception as exc:
                stats['errors'] += 1
                tqdm.write(f"{c.FAIL}{c.BOLD}[致命错] {rel_path}: 获取任务结果失败 - {exc}{c.ENDC}")
                logging.critical(f"获取任务结果失败 [{rel_path}]:", exc_info=True)
            progress_bar.update(1)
        progress_bar.close()
    duration = time.time() - start_time
    print_summary(stats, duration, log_file, c)

if __name__ == "__main__":
    main()
